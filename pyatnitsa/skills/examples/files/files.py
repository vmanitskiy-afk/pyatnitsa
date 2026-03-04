"""
Навык files — работа с файлами и папками пользователя.
Workspace настраивается через FILES_WORKSPACE env или settings.
"""
from __future__ import annotations

import base64
import csv
import io
import json
import mimetypes
import os
import shutil
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import structlog

from pyatnitsa.skills.base import BaseSkill, LLMTool

logger = structlog.get_logger("skill.files")

MAX_READ_SIZE = 10 * 1024 * 1024   # 10 MB
MAX_EXCEL_ROWS = 10_000
MAX_SEARCH_RESULTS = 100


class FileSkill(BaseSkill):
    name = "files"
    description = "Работа с файлами: чтение, анализ, поиск, создание (xlsx, csv, pdf, docx, txt, изображения)"
    version = "1.0.0"

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self._workspace: Path | None = None

    async def on_load(self):
        ws = os.environ.get("FILES_WORKSPACE", "")
        if ws:
            self._workspace = Path(ws).resolve()
            if not self._workspace.exists():
                self._workspace.mkdir(parents=True, exist_ok=True)
            logger.info("files_skill_loaded", workspace=str(self._workspace))
        else:
            logger.warning("files_skill_no_workspace",
                           hint="Установите FILES_WORKSPACE в .env или через админку")

    def _safe_path(self, user_path: str) -> Path:
        """Резолвит путь внутри workspace, блокирует traversal."""
        if not self._workspace:
            raise RuntimeError("Рабочая папка не настроена. Укажите FILES_WORKSPACE в настройках.")
        clean = user_path.replace("\\", "/").lstrip("/")
        target = (self._workspace / clean).resolve()
        if not str(target).startswith(str(self._workspace)):
            raise PermissionError(f"Доступ запрещён: выход за пределы рабочей папки")
        return target

    # ─── Tools ───────────────────────────────────────────

    def get_tools(self) -> list[LLMTool]:
        return [
            # Navigation
            LLMTool("files.list", "Список файлов и папок в директории", {
                "type": "object", "properties": {
                    "path": {"type": "string", "description": "Путь относительно workspace (по умолчанию корень)", "default": ""},
                    "recursive": {"type": "boolean", "description": "Рекурсивно", "default": False},
                    "pattern": {"type": "string", "description": "Glob-паттерн фильтра (напр. *.xlsx)", "default": "*"},
                },
            }),
            LLMTool("files.tree", "Дерево структуры папки", {
                "type": "object", "properties": {
                    "path": {"type": "string", "default": ""},
                    "depth": {"type": "integer", "description": "Глубина (1-5)", "default": 2},
                },
            }),
            LLMTool("files.stats", "Сводка по папке: типы файлов, размер, последние изменения", {
                "type": "object", "properties": {
                    "path": {"type": "string", "default": ""},
                },
            }),
            # Reading
            LLMTool("files.read", "Прочитать текстовый файл (txt, csv, json, xml, md, py, js...)", {
                "type": "object", "properties": {
                    "path": {"type": "string", "description": "Путь к файлу"},
                    "encoding": {"type": "string", "default": "auto"},
                    "lines": {"type": "integer", "description": "Макс. строк (0 = все)", "default": 0},
                }, "required": ["path"],
            }),
            LLMTool("files.read_excel", "Прочитать данные из Excel файла (xlsx/xls)", {
                "type": "object", "properties": {
                    "path": {"type": "string"},
                    "sheet": {"type": "string", "description": "Имя листа (по умолчанию первый)"},
                    "max_rows": {"type": "integer", "default": 100},
                }, "required": ["path"],
            }),
            LLMTool("files.read_pdf", "Извлечь текст из PDF", {
                "type": "object", "properties": {
                    "path": {"type": "string"},
                    "pages": {"type": "string", "description": "Номера страниц через запятую (по умолчанию все)"},
                }, "required": ["path"],
            }),
            LLMTool("files.read_docx", "Извлечь текст из Word-документа", {
                "type": "object", "properties": {
                    "path": {"type": "string"},
                }, "required": ["path"],
            }),
            LLMTool("files.read_image", "Прочитать изображение (возвращает base64 для анализа)", {
                "type": "object", "properties": {
                    "path": {"type": "string"},
                }, "required": ["path"],
            }),
            # Search
            LLMTool("files.search", "Поиск текста в файлах (grep)", {
                "type": "object", "properties": {
                    "query": {"type": "string", "description": "Текст для поиска"},
                    "path": {"type": "string", "default": ""},
                    "file_types": {"type": "string", "description": "Расширения через запятую (напр. txt,csv,py)"},
                    "case_sensitive": {"type": "boolean", "default": False},
                }, "required": ["query"],
            }),
            # Writing
            LLMTool("files.write", "Создать или перезаписать текстовый файл", {
                "type": "object", "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string"},
                    "encoding": {"type": "string", "default": "utf-8"},
                }, "required": ["path", "content"],
            }),
            LLMTool("files.write_csv", "Создать CSV файл из данных", {
                "type": "object", "properties": {
                    "path": {"type": "string"},
                    "headers": {"type": "array", "items": {"type": "string"}},
                    "rows": {"type": "array", "items": {"type": "array"}},
                    "delimiter": {"type": "string", "default": ","},
                }, "required": ["path", "headers", "rows"],
            }),
            LLMTool("files.mkdir", "Создать директорию", {
                "type": "object", "properties": {
                    "path": {"type": "string"},
                }, "required": ["path"],
            }),
            LLMTool("files.copy", "Копировать файл", {
                "type": "object", "properties": {
                    "src": {"type": "string"},
                    "dst": {"type": "string"},
                }, "required": ["src", "dst"],
            }),
            LLMTool("files.move", "Переместить или переименовать файл", {
                "type": "object", "properties": {
                    "src": {"type": "string"},
                    "dst": {"type": "string"},
                }, "required": ["src", "dst"],
            }),
        ]

    async def execute(self, tool_name: str, params: dict) -> str:
        cmd = tool_name.split(".")[-1]
        dispatch = {
            "list": self._list, "tree": self._tree, "stats": self._stats,
            "read": self._read, "read_excel": self._read_excel,
            "read_pdf": self._read_pdf, "read_docx": self._read_docx,
            "read_image": self._read_image,
            "search": self._search,
            "write": self._write, "write_csv": self._write_csv,
            "mkdir": self._mkdir, "copy": self._copy, "move": self._move,
        }
        fn = dispatch.get(cmd)
        if not fn:
            return json.dumps({"error": f"Неизвестная команда: {tool_name}"}, ensure_ascii=False)
        try:
            return await fn(params)
        except (PermissionError, RuntimeError) as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)
        except Exception as e:
            logger.error("files_skill_error", cmd=cmd, error=str(e)[:200])
            return json.dumps({"error": f"Ошибка: {e}"}, ensure_ascii=False)

    # ─── Navigation ──────────────────────────────────────

    async def _list(self, p: dict) -> str:
        target = self._safe_path(p.get("path", ""))
        if not target.exists():
            return json.dumps({"error": f"Путь не найден: {p.get('path', '')}"}, ensure_ascii=False)

        recursive = p.get("recursive", False)
        pattern = p.get("pattern", "*")
        items = []

        if target.is_file():
            items.append(self._file_info(target))
        else:
            glob_fn = target.rglob if recursive else target.glob
            for f in sorted(glob_fn(pattern)):
                if f.name.startswith("."):
                    continue
                items.append(self._file_info(f))
                if len(items) >= 200:
                    break

        return json.dumps({"path": str(target.relative_to(self._workspace)),
                           "count": len(items), "items": items}, ensure_ascii=False)

    async def _tree(self, p: dict) -> str:
        target = self._safe_path(p.get("path", ""))
        depth = min(p.get("depth", 2), 5)
        lines = []
        self._build_tree(target, "", depth, lines)
        return "\n".join(lines) if lines else "Пустая директория"

    async def _stats(self, p: dict) -> str:
        target = self._safe_path(p.get("path", ""))
        if not target.is_dir():
            return json.dumps({"error": "Путь не является директорией"}, ensure_ascii=False)

        by_ext: dict[str, dict] = {}
        total_size = 0
        total_files = 0
        latest_file = ""
        latest_mtime = 0

        for f in target.rglob("*"):
            if f.is_file() and not f.name.startswith("."):
                ext = f.suffix.lower() or "(без расширения)"
                size = f.stat().st_size
                mtime = f.stat().st_mtime
                total_size += size
                total_files += 1
                if ext not in by_ext:
                    by_ext[ext] = {"count": 0, "size": 0}
                by_ext[ext]["count"] += 1
                by_ext[ext]["size"] += size
                if mtime > latest_mtime:
                    latest_mtime = mtime
                    latest_file = str(f.relative_to(self._workspace))

        return json.dumps({
            "path": str(target.relative_to(self._workspace)),
            "total_files": total_files,
            "total_size": self._fmt_size(total_size),
            "total_size_bytes": total_size,
            "by_extension": dict(sorted(by_ext.items(), key=lambda x: -x[1]["count"])),
            "latest_modified": latest_file,
            "latest_modified_time": datetime.fromtimestamp(latest_mtime).isoformat() if latest_mtime else None,
        }, ensure_ascii=False)

    # ─── Reading ─────────────────────────────────────────

    async def _read(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        if not target.is_file():
            return json.dumps({"error": f"Файл не найден: {p['path']}"}, ensure_ascii=False)
        if target.stat().st_size > MAX_READ_SIZE:
            return json.dumps({"error": f"Файл слишком большой ({self._fmt_size(target.stat().st_size)})"}, ensure_ascii=False)

        encoding = p.get("encoding", "auto")
        if encoding == "auto":
            encoding = self._detect_encoding(target)

        text = target.read_text(encoding=encoding, errors="replace")
        max_lines = p.get("lines", 0)
        if max_lines > 0:
            lines = text.splitlines()
            text = "\n".join(lines[:max_lines])
            if len(lines) > max_lines:
                text += f"\n\n... (ещё {len(lines) - max_lines} строк)"

        return text

    async def _read_excel(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        if not target.is_file():
            return json.dumps({"error": f"Файл не найден: {p['path']}"}, ensure_ascii=False)

        try:
            import openpyxl
        except ImportError:
            return json.dumps({"error": "openpyxl не установлен (pip install openpyxl)"}, ensure_ascii=False)

        wb = openpyxl.load_workbook(str(target), read_only=True, data_only=True)
        sheet_name = p.get("sheet") or wb.sheetnames[0]
        if sheet_name not in wb.sheetnames:
            return json.dumps({"error": f"Лист '{sheet_name}' не найден. Доступные: {wb.sheetnames}"}, ensure_ascii=False)

        ws = wb[sheet_name]
        max_rows = min(p.get("max_rows", 100), MAX_EXCEL_ROWS)
        rows = []
        headers = []

        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c) if c is not None else "" for c in row]
            else:
                rows.append([self._cell_value(c) for c in row])
            if i >= max_rows:
                break

        wb.close()
        total = ws.max_row - 1 if ws.max_row else 0

        return json.dumps({
            "file": p["path"],
            "sheet": sheet_name,
            "sheets": wb.sheetnames,
            "headers": headers,
            "rows": rows,
            "rows_shown": len(rows),
            "total_rows": total,
        }, ensure_ascii=False)

    async def _read_pdf(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        if not target.is_file():
            return json.dumps({"error": f"Файл не найден: {p['path']}"}, ensure_ascii=False)

        try:
            import fitz  # PyMuPDF
        except ImportError:
            return json.dumps({"error": "PyMuPDF не установлен (pip install PyMuPDF)"}, ensure_ascii=False)

        doc = fitz.open(str(target))
        pages_str = p.get("pages", "")
        if pages_str:
            page_nums = [int(x.strip()) - 1 for x in pages_str.split(",") if x.strip().isdigit()]
        else:
            page_nums = list(range(len(doc)))

        parts = []
        for pn in page_nums:
            if 0 <= pn < len(doc):
                text = doc[pn].get_text()
                if text.strip():
                    parts.append(f"--- Страница {pn + 1} ---\n{text}")

        doc.close()
        if not parts:
            return "PDF не содержит извлекаемого текста (возможно, сканированный документ)."
        return "\n\n".join(parts)

    async def _read_docx(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        if not target.is_file():
            return json.dumps({"error": f"Файл не найден: {p['path']}"}, ensure_ascii=False)

        try:
            import docx
        except ImportError:
            return json.dumps({"error": "python-docx не установлен (pip install python-docx)"}, ensure_ascii=False)

        doc = docx.Document(str(target))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Таблицы
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append(f"\n[Таблица {i + 1}]\n" + "\n".join(rows))

        return "\n".join(parts) if parts else "Документ пуст."

    async def _read_image(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        if not target.is_file():
            return json.dumps({"error": f"Файл не найден: {p['path']}"}, ensure_ascii=False)
        if target.stat().st_size > MAX_READ_SIZE:
            return json.dumps({"error": "Изображение слишком большое"}, ensure_ascii=False)

        mime = mimetypes.guess_type(str(target))[0] or "image/jpeg"
        data = target.read_bytes()
        b64 = base64.b64encode(data).decode("ascii")
        return json.dumps({
            "type": "image",
            "mime": mime,
            "size": len(data),
            "base64": b64[:100] + "...",  # Для лога; полный base64 отдаётся через vision
            "path": p["path"],
        }, ensure_ascii=False)

    # ─── Search ──────────────────────────────────────────

    async def _search(self, p: dict) -> str:
        query = p["query"]
        target = self._safe_path(p.get("path", ""))
        case_sensitive = p.get("case_sensitive", False)
        file_types = p.get("file_types", "")
        extensions = {f".{e.strip().lstrip('.')}" for e in file_types.split(",") if e.strip()} if file_types else None

        if not case_sensitive:
            query_lower = query.lower()

        results = []
        text_exts = {".txt", ".csv", ".json", ".xml", ".md", ".py", ".js", ".ts",
                     ".html", ".css", ".yml", ".yaml", ".toml", ".ini", ".cfg",
                     ".log", ".sql", ".sh", ".bat", ".env", ".conf"}

        search_exts = extensions or text_exts

        for f in target.rglob("*"):
            if not f.is_file() or f.name.startswith("."):
                continue
            if f.suffix.lower() not in search_exts:
                continue
            if f.stat().st_size > 5 * 1024 * 1024:  # 5MB max per file
                continue

            try:
                content = f.read_text(encoding="utf-8", errors="replace")
                for line_num, line in enumerate(content.splitlines(), 1):
                    match = (query in line) if case_sensitive else (query_lower in line.lower())
                    if match:
                        results.append({
                            "file": str(f.relative_to(self._workspace)),
                            "line": line_num,
                            "text": line.strip()[:200],
                        })
                        if len(results) >= MAX_SEARCH_RESULTS:
                            return json.dumps({"query": query, "results": results,
                                               "truncated": True}, ensure_ascii=False)
            except Exception:
                continue

        return json.dumps({"query": query, "results": results,
                           "count": len(results)}, ensure_ascii=False)

    # ─── Writing ─────────────────────────────────────────

    async def _write(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        target.parent.mkdir(parents=True, exist_ok=True)
        encoding = p.get("encoding", "utf-8")
        target.write_text(p["content"], encoding=encoding)
        return json.dumps({"success": True, "path": p["path"],
                           "size": self._fmt_size(target.stat().st_size)}, ensure_ascii=False)

    async def _write_csv(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        target.parent.mkdir(parents=True, exist_ok=True)
        delimiter = p.get("delimiter", ",")

        with open(target, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f, delimiter=delimiter)
            writer.writerow(p["headers"])
            writer.writerows(p["rows"])

        return json.dumps({"success": True, "path": p["path"],
                           "rows": len(p["rows"]),
                           "size": self._fmt_size(target.stat().st_size)}, ensure_ascii=False)

    async def _mkdir(self, p: dict) -> str:
        target = self._safe_path(p["path"])
        target.mkdir(parents=True, exist_ok=True)
        return json.dumps({"success": True, "path": p["path"]}, ensure_ascii=False)

    async def _copy(self, p: dict) -> str:
        src = self._safe_path(p["src"])
        dst = self._safe_path(p["dst"])
        if not src.exists():
            return json.dumps({"error": f"Источник не найден: {p['src']}"}, ensure_ascii=False)
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(src), str(dst))
        return json.dumps({"success": True, "src": p["src"], "dst": p["dst"]}, ensure_ascii=False)

    async def _move(self, p: dict) -> str:
        src = self._safe_path(p["src"])
        dst = self._safe_path(p["dst"])
        if not src.exists():
            return json.dumps({"error": f"Источник не найден: {p['src']}"}, ensure_ascii=False)
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dst))
        return json.dumps({"success": True, "src": p["src"], "dst": p["dst"]}, ensure_ascii=False)

    # ─── Helpers ─────────────────────────────────────────

    def _file_info(self, path: Path) -> dict:
        """Метаданные файла/папки."""
        st = path.stat()
        return {
            "name": path.name,
            "path": str(path.relative_to(self._workspace)),
            "type": "dir" if path.is_dir() else "file",
            "size": st.st_size if path.is_file() else None,
            "size_human": self._fmt_size(st.st_size) if path.is_file() else None,
            "modified": datetime.fromtimestamp(st.st_mtime).isoformat(),
            "extension": path.suffix.lower() if path.is_file() else None,
        }

    def _build_tree(self, path: Path, prefix: str, depth: int, lines: list):
        """Рекурсивно строит дерево."""
        if depth < 0:
            return
        try:
            entries = sorted(path.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
        except PermissionError:
            return
        entries = [e for e in entries if not e.name.startswith(".")]
        for i, entry in enumerate(entries):
            is_last = i == len(entries) - 1
            connector = "└── " if is_last else "├── "
            size_info = f" ({self._fmt_size(entry.stat().st_size)})" if entry.is_file() else ""
            lines.append(f"{prefix}{connector}{entry.name}{size_info}")
            if entry.is_dir() and depth > 0:
                extension = "    " if is_last else "│   "
                self._build_tree(entry, prefix + extension, depth - 1, lines)

    @staticmethod
    def _fmt_size(size: int) -> str:
        for unit in ("Б", "КБ", "МБ", "ГБ"):
            if size < 1024:
                return f"{size:.0f} {unit}" if unit == "Б" else f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} ТБ"

    @staticmethod
    def _detect_encoding(path: Path) -> str:
        """Определяет кодировку файла."""
        try:
            import chardet
            raw = path.read_bytes()[:8192]
            result = chardet.detect(raw)
            return result.get("encoding", "utf-8") or "utf-8"
        except ImportError:
            return "utf-8"

    @staticmethod
    def _cell_value(val: Any) -> Any:
        """Сериализует значение ячейки Excel."""
        if val is None:
            return ""
        if isinstance(val, datetime):
            return val.isoformat()
        return val
