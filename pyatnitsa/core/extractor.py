"""Извлечение текста из файлов для LLM-анализа."""

from __future__ import annotations

import mimetypes
from pathlib import Path

import structlog

logger = structlog.get_logger()

MAX_TEXT_LEN = 30_000  # ~10K токенов, безопасно для контекста


async def extract_text(file_path: str, mime_type: str | None = None) -> str | None:
    """Извлекает текст из файла. Возвращает None если не удалось."""
    path = Path(file_path)
    if not path.exists():
        return None

    if not mime_type:
        mime_type = mimetypes.guess_type(str(path))[0] or ""

    try:
        if mime_type == "application/pdf":
            return _extract_pdf(path)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return _extract_docx(path)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ):
            return _extract_xlsx(path)
        elif mime_type.startswith("text/") or mime_type in ("application/json", "application/xml"):
            return _extract_text(path)
        elif mime_type.startswith("image/"):
            return None  # Images handled separately if vision is available
        else:
            return None
    except Exception as e:
        logger.warning("extract_text_error", path=str(path), error=str(e))
        return None


def _extract_pdf(path: Path) -> str | None:
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text.strip())
        doc.close()
        result = "\n\n".join(pages)
        return result[:MAX_TEXT_LEN] if result else None
    except ImportError:
        logger.warning("pymupdf_not_installed")
        return None


def _extract_docx(path: Path) -> str | None:
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        result = "\n".join(paragraphs)
        return result[:MAX_TEXT_LEN] if result else None
    except ImportError:
        logger.warning("python_docx_not_installed")
        return None


def _extract_xlsx(path: Path) -> str | None:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"--- Лист: {ws.title} ---")
            for row in ws.iter_rows(max_row=200, values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    lines.append(" | ".join(cells))
        wb.close()
        result = "\n".join(lines)
        return result[:MAX_TEXT_LEN] if result else None
    except ImportError:
        logger.warning("openpyxl_not_installed")
        return None


def _extract_text(path: Path) -> str | None:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return text[:MAX_TEXT_LEN] if text.strip() else None
    except Exception:
        return None
